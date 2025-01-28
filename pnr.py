import re
import json 
import os
from tkinter import messagebox, filedialog
import tkinter as tk
from tkinter import ttk
from tkinter import Menu
import docx
# type: ignore
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create the main window
root = tk.Tk()
root.title("Patch Note Releaser")
root.geometry("950x880")
root.minsize(800, 600)

# Variables to hold form data
editor_mode_data = {
    "project_name": "",
    "folder_path": "Выберите папку для сохранения проекта",
    "main_text": "",
    "version": "",
    "date": "",
    "description": "",
    "improvements": "",
    "fixed_bugs": "",
    "notes": ""
}

release_mode_data = {
    "folder_path": "Выберите папку для сохранения проекта",
    "project_name": "",
    "version": "",
    "date": "",
    "title": "",
    "summary": "",
    "summary_checked": False,
    "sections": {
        "New Features": {"entries": [], "checked": False},
        "Improvements": {"entries": [], "checked": False},
        "Bug Fixes": {"entries": [], "checked": False},
        "Known Issues": {"entries": [], "checked": False}
    }
}

# Переменные для отслеживания текущего режима
current_mode = None  # Изменяем с "editor" на None

SETTINGS_FILE = "settings.json"

# Functions

def main():
    """Инициализация программы"""
    global current_mode
    
    # Загружаем сохраненный режим или используем release по умолчанию
    default_mode = get_setting('default_mode', 'release')
    current_mode = default_mode  # Устанавливаем текущий режим
    
    if default_mode == 'editor':
        switch_to_editor_mode()
    else:
        switch_to_release_mode()


def choose_folder(var):
    """Открывает диалог выбора папки и сохраняет путь в переданную переменную."""
    folder_path = filedialog.askdirectory()
    if folder_path:
        var.set(folder_path)

def choose_folder(path_var):
    """Открывает диалог выбора папки и сохраняет путь в переменную."""
    folder_path = filedialog.askdirectory()
    if folder_path:  # Проверяем, что путь был выбран
        path_var.set(folder_path)

def add_field(frame):
    """Добавляет новое текстовое поле в указанный фрейм."""
    entry = tk.Entry(frame, width=80)
    entry.pack(pady=2, padx=10, fill=tk.X)

def remove_field(frame):
    if len(frame.winfo_children()) > 1:
        frame.winfo_children()[-1].destroy()


def exit_program():
    """Завершает выполнение программы."""
    root.destroy()

# Functions for mode switching and actions
def switch_to_editor_mode():
    """Переключает интерфейс в режим редактора."""
    global current_mode
    # Сохраняем текущие данные перед переключением
    try:
        save_form_data(force_save=True)
    except:
        print("Не удалось сохранить текущие данные")
    
    clear_main_frame()
    current_mode = "editor"
    editor_mode_ui()
    load_form_data()

def switch_to_release_mode():
    """Переключает интерфейс в режим релиза."""
    global current_mode
    # Сохраняем текущие данные перед переключением
    try:
        save_form_data(force_save=True)
    except:
        print("Не удалось сохранить текущие данные")
    
    clear_main_frame()
    current_mode = "release"
    release_mode_ui()
    load_form_data()

def clear_main_frame():
    """Очищает основной фрейм от всех виджетов."""
    for widget in main_frame.winfo_children():
        widget.destroy()

def save_current_data():
    """Сохраняет текущие данные формы в соответствующий словарь."""
    if not main_frame.winfo_children():
        return
        
    if current_mode == "editor":
        # Save Editor Mode data
        editor_mode_data.update({
            "project_name": form1_entry.get(),
            "folder_path": folder_path_var.get(),
            "main_text": form2_text.get("1.0", tk.END).strip(),
            "version": version_entry.get(),
            "date": date_entry.get(),
            "description": description_entry.get(),
            "improvements": form4_text.get("1.0", tk.END).strip(),
            "fixed_bugs": form6_text.get("1.0", tk.END).strip(),
            "notes": form7_text.get("1.0", tk.END).strip(),
            "form6_checked": form6_check_var.get(),
            "form7_checked": form7_check_var.get()
        })
    else:
        # Save Release Mode data
        try:
            release_mode_data.update({
                "folder_path": release_folder_path_var.get(),
                "project_name": release_project_entry.get() if hasattr(release_project_entry, 'get') else "",
                "version": version_entry.get() if hasattr(version_entry, 'get') else "",
                "date": date_entry.get() if hasattr(date_entry, 'get') else "",
                "title": title_entry.get() if hasattr(title_entry, 'get') else "",
                "summary": summary_text.get("1.0", tk.END).strip() if 'summary_text' in globals() else "",
                "summary_checked": summary_check_var.get() if 'summary_check_var' in globals() else False
            })
            
            # Сохраняем секции
            if 'section_widgets' in globals():
                for section, widgets_data in section_widgets.items():
                    entries = []
                    if "entries" in widgets_data:
                        for entry in widgets_data["entries"]:
                            if isinstance(entry, tk.Entry):
                                entries.append(entry.get())
                    
                    is_checked = False
                    if "check_var" in widgets_data:
                        is_checked = widgets_data["check_var"].get()
                    
                    release_mode_data["sections"][section] = {
                        "entries": entries if any(entries) else [""],
                        "checked": is_checked
                    }

        except (AttributeError, tk.TclError) as e:
            print(f"Warning: Some widgets were not accessible: {e}")

def create_text_widget_with_menu(parent, height=10):
    """Создает текстовое поле с контекстным меню."""
    text_widget = tk.Text(parent, wrap=tk.WORD, height=height)
    
    # Создаем контекстное меню
    context_menu = Menu(text_widget, tearoff=0)
    context_menu.add_command(label="Копировать", command=lambda: text_widget.event_generate("<<Copy>>"))
    context_menu.add_command(label="Вставить", command=lambda: text_widget.event_generate("<<Paste>>"))
    context_menu.add_command(label="Вырезать", command=lambda: text_widget.event_generate("<<Cut>>"))
    context_menu.add_separator()
    context_menu.add_command(label="Выделить всё", command=lambda: text_widget.tag_add("sel", "1.0", "end"))
    
    # Привязываем появление меню к правой кнопке мыши
    def show_menu(event):
        context_menu.tk_popup(event.x_root, event.y_root)
    
    text_widget.bind("<Button-3>", show_menu)
    return text_widget

def create_text_widget_with_menu_and_format(parent, height=10):
    """Создает текстовое поле с контекстным меню и автоформатированием."""
    text_widget = tk.Text(parent, wrap=tk.WORD, height=height)
    scrollbar = tk.Scrollbar(parent, command=text_widget.yview)
    text_widget.configure(yscrollcommand=scrollbar.set)
    
    # Создаем контекстное меню
    context_menu = Menu(text_widget, tearoff=0)
    context_menu.add_command(label="Копировать", command=lambda: text_widget.event_generate("<<Copy>>"))
    context_menu.add_command(label="Вставить", command=lambda: text_widget.event_generate("<<Paste>>"))
    context_menu.add_command(label="Вырезать", command=lambda: text_widget.event_generate("<<Cut>>"))
    context_menu.add_separator()
    context_menu.add_command(label="Выделить всё", command=lambda: text_widget.tag_add("sel", "1.0", "end"))
    
    def show_menu(event):
        context_menu.tk_popup(event.x_root, event.y_root)
    
    def format_text(event):
        # Получаем текущую позицию курсора
        current_pos = text_widget.index("insert")
        line_num = int(current_pos.split('.')[0])
        
        # Получаем текст текущей строки
        line_start = f"{line_num}.0"
        line_end = f"{line_num}.end"
        current_line = text_widget.get(line_start, line_end).strip()
        
        # Форматируем текущую строку если она не пустая
        if current_line:
            if not current_line.startswith('- '):
                current_line = f"- {current_line}"
            if not current_line.endswith(';'):
                current_line = f"{current_line};"
            
            # Заменяем текущую строку на отформатированную
            text_widget.delete(line_start, line_end)
            text_widget.insert(line_start, current_line)
        
        # Добавляем новую строку
        text_widget.insert("insert", "\n")
        return "break"  # Предотвращает стандартное поведение Enter
    
    text_widget.bind("<Button-3>", show_menu)
    text_widget.bind("<Return>", format_text)
    return text_widget, scrollbar

def create_section_with_buttons(parent, section_name, initial_entries=1):
    """Создает секцию с кнопками добавления/удаления полей."""
    frame = tk.Frame(parent)
    entries_frame = tk.Frame(frame)
    entries = []
    
    def add_entry():
        entry = tk.Entry(entries_frame, width=80)
        entry.pack(padx=5, pady=2)
        entries.append(entry)
        
    def remove_entry():
        if len(entries) > 1:  # Всегда оставляем хотя бы одно поле
            entry = entries.pop()
            entry.destroy()
    
    # Создаем чекбокс
    check_var = tk.BooleanVar(value=False)
    check = tk.Checkbutton(frame, text=section_name, variable=check_var)
    check.var = check_var
    check.pack(anchor="w", padx=5)
    
    # Создаем начальные поля ввода
    for _ in range(initial_entries):
        add_entry()
    
    entries_frame.pack(fill=tk.X)

  # Кнопки управления
    buttons_frame = tk.Frame(frame)
    add_btn = ttk.Button(buttons_frame, text="+", width=3, command=add_entry)
    remove_btn = ttk.Button(buttons_frame, text="-", width=3, command=remove_entry)
    
    add_btn.pack(side=tk.LEFT, padx=2)
    remove_btn.pack(side=tk.LEFT, padx=2)
    buttons_frame.pack(anchor="w", padx=5, pady=2)
    
    return frame, entries, check_var

def validate_version(P):
    """Проверяет ввод версии в формате XX.XX.XXX"""
    if P == "":
        return True
    parts = P.split('.')
    if len(parts) > 3:
        return False
    try:
        for i, part in enumerate(parts):
            if not part.isdigit() and part != "":
                return False
            if part != "":
                if i == 0 and len(part) > 2:  # Major Version (XX)
                    return False
                if i == 1 and len(part) > 2:  # Minor Version (XX)
                    return False
                if i == 2 and len(part) > 3:  # Patch Version (XXX)
                    return False
        return True
    except ValueError:
        return False


def format_version(event):
    """Форматирует версию, сохраняя ведущие нули."""
    entry = event.widget
    version = entry.get()
    
    if version:
        parts = version.split('.')
        formatted_parts = []
        for part in parts:
            if part:
                # Сохраняем ведущие нули
                formatted_parts.append(part)
            else:
                formatted_parts.append(part)
        
        formatted_version = '.'.join(formatted_parts)
        if formatted_version != version:
            entry.delete(0, tk.END)
            entry.insert(0, formatted_version)

def validate_date(P):
    """Проверяет ввод даты в формате ДД.ММ.ГГГГ"""
    if P == "":
        return True
    
    # Разрешаем ввод только цифр и точек
    if not all(c.isdigit() or c == '.' for c in P):
        return False
    
    parts = P.split('.')
    if len(parts) > 3:
        return False
    
    try:
        # Проверяем длину каждой части
        if len(parts) >= 1 and len(parts[0]) > 2:  # ДД
            return False
        if len(parts) >= 2 and len(parts[1]) > 2:  # ММ
            return False
        if len(parts) >= 3 and len(parts[2]) > 4:  # ГГГГ
            return False
            
        # Проверяем валидность значений только для полностью заполненных частей
        if len(parts[0]) == 2 and int(parts[0]) > 0:
            if not (1 <= int(parts[0]) <= 31):
                return False
        if len(parts) > 1 and len(parts[1]) == 2 and int(parts[1]) > 0:
            if not (1 <= int(parts[1]) <= 12):
                return False
        if len(parts) > 2 and len(parts[2]) == 4:
            if not (1000 <= int(parts[2]) <= 9999):
                return False
        
        return True
    except (ValueError, IndexError):
        return True

def format_date(event):
    """Форматирует дату, автоматически добавляя точки."""
    entry = event.widget
    date = entry.get()
    
    # Если пользователь вводит точку, пропускаем форматирование
    if event.char == '.':
        return
    
    # Убираем все точки для обработки
    clean_date = date.replace('.', '')
    formatted = ''
    
    # Форматируем дату
    if len(clean_date) > 0:
        formatted += clean_date[:2]
        if len(clean_date) > 2:
            formatted += '.' + clean_date[2:4]
            if len(clean_date) > 4:
                formatted += '.' + clean_date[4:8]
    
    # Обновляем поле только если формат изменился
    if formatted != date:
        entry.delete(0, tk.END)
        entry.insert(0, formatted)

def clear_frame(frame):
    """Очистка всех виджетов из переданного фрейма."""
    for widget in frame.winfo_children():
        widget.destroy()

def save_to_file():
    """Сохраняет данные текущего режима в файл."""
    save_current_data()  # Сначала сохраняем текущие данные

    # Определяем SEPARATOR в начале функции
    SEPARATOR = "=" * 70 if current_mode == "editor" else "_" * 70

    try:
        # Получаем базовую информацию в зависимости от режима
        if current_mode == "editor":
            folder_path = folder_path_var.get()
            if folder_path == "Выберите папку для сохранения проекта":
                messagebox.showerror("Ошибка", "Выберите папку для сохранения проекта.")
                return
            
            project_name = form1_entry.get()
            if not project_name:
                messagebox.showerror("Ошибка", "Введите название проекта.")
                return
            
            version = version_entry.get()
            version_at_end = get_setting('version_at_end', True)
        else:  # release mode
            project_name = release_project_entry.get().strip()
            if not project_name:
                messagebox.showerror("Ошибка", "Введите название проекта")
                return
                
            folder_path = release_folder_path_var.get()
            if not folder_path or folder_path == "Выберите папку для сохранения проекта":
                messagebox.showerror("Ошибка", "Выберите папку для сохранения")
                return
            
            version = version_entry.get()
            version_at_end = get_setting('release_version_at_end', True)

        # Формируем базовое имя файла
        if version_at_end:
            base_name = f"{project_name}_v{version}"
        else:
            base_name = f"v{version}_{project_name}"

        # Получаем выбранный формат файла из настроек
        file_format = get_setting('default_file_format', '.txt')
        file_path = os.path.join(folder_path, base_name + file_format)

        # Получаем контент в зависимости от режима
        if current_mode == "editor":
            content = create_editor_content(SEPARATOR)
        else:
            content = create_release_content(SEPARATOR)

        # Сохраняем файл в выбранном формате
        if file_format == '.txt':
            save_as_txt(file_path, content)
        elif file_format == '.docx':
            save_as_docx(file_path, content)
        elif file_format == '.rtf':
            save_as_rtf(file_path, content)
        elif file_format == '.pdf':
            save_as_pdf(file_path, content)

        messagebox.showinfo("Успех", f"Файл успешно сохранен:\n{file_path}")
        
    except Exception as e:
        messagebox.showerror("Ошибка", f"Ошибка при сохранении файла:\n{str(e)}")

def create_section_header(title, separator_char='='):
    """Создает заголовок секции с равномерным распределением символов разделителя.
    
    Args:
        title (str): Название секции
        separator_char (str): Символ разделителя ('=' для Editor Mode, '_' для Release Mode)
    """
    total_length = 70  # Увеличиваем длину с 50 до 70
    title_length = len(title)
    remaining_space = total_length - title_length - 2  # -2 для пробелов вокруг заголовка
    left_padding = remaining_space // 2
    right_padding = remaining_space - left_padding
    return f"{separator_char * left_padding} {title} {separator_char * right_padding}"

def create_editor_content(SEPARATOR):
    """Создает контент для режима Editor"""
    content = [
        SEPARATOR,
        f"Версия правки - {version_entry.get()} | Дата - {date_entry.get()} | Описание - {description_entry.get()}",
        SEPARATOR,
        create_section_header("Улучшения:", "="),
        ""
    ]
    
    # Обрабатываем текст улучшений
    improvements = form4_text.get('1.0', tk.END).strip()
    if improvements:
        # Разбиваем по точке с запятой и добавляем каждый пункт с новой строки
        items = improvements.split(';')
        for item in items:
            if item.strip():
                if not item.strip().startswith('-'):
                    item = '- ' + item.strip()
                content.append(item)
    
    content.extend([
        "",
        SEPARATOR,
        create_section_header("Основное текстовое поле:", "="),
        "",
        form2_text.get('1.0', tk.END).strip(),
        "",
        SEPARATOR
    ])
    
    # Добавляем опциональные секции если отмечены чекбоксы
    if form6_check_var.get():  # Fixed Bugs
        bugs_text = form6_text.get('1.0', tk.END).strip()
        if bugs_text:
            content.append(create_section_header("Исправленные ошибки:", "="))
            content.append("")
            # Разбиваем текст ошибок на отдельные пункты
            for bug in bugs_text.split(';'):
                if bug.strip():
                    if not bug.strip().startswith('-'):
                        bug = '- ' + bug.strip()
                    content.append(bug)
            content.append("")
            content.append(SEPARATOR)
            
    if form7_check_var.get():  # Notes
        notes_text = form7_text.get('1.0', tk.END).strip()
        if notes_text:
            content.append(create_section_header("Заметки:", "="))
            content.append("")
            # Разбиваем текст заметок на отдельные пункты
            for note in notes_text.split(';'):
                if note.strip():
                    if not note.strip().startswith('-'):
                        note = '- ' + note.strip()
                    content.append(note)
            content.append("")
            content.append(SEPARATOR)
    
    return content

def create_release_content(SEPARATOR):
    """Создает контент для режима Release"""
    content = ["\n"]
    
    # Добавляем базовую информацию
    if version_entry.get().strip():
        content.append(f"Версия: {version_entry.get().strip()}")
    if date_entry.get().strip():
        content.append(f"Дата: {date_entry.get().strip()}")
    if title_entry.get().strip():
        content.append(f"Описание: {title_entry.get().strip()}")  # Убираем \n
    
    content.append("")  # Одна пустая строка перед разделителем
    content.append(SEPARATOR)
    
    # Добавляем Summary если он отмечен
    if summary_check_var.get():
        summary_content = summary_text.get("1.0", tk.END).strip()
        if summary_content:
            content.append(create_section_header("Общее краткое содержание\\резюме:", "_"))
            content.append("")
            content.append(summary_content)
            content.append("")
            content.append("_" * 70)
    
    # Словарь для перевода названий секций
    sections_translation = {
        "Summary": "Общее краткое содержание\\резюме:",
        "New Features": "Новые функции:",
        "Improvements": "Улучшения:",
        "Bug Fixes": "Исправление ошибок:",
        "Known Issues": "Известные проблемы:",
        "Notes": "Заметки:",
        "Feedback": "Обратная связь:"
    }
    
    
    # Добавляем остальные секции
    for section, widgets_data in section_widgets.items():
        if widgets_data["check_var"].get():
            section_content = []
            for entry in widgets_data["entries"]:
                text = entry.get().strip()
                if text:
                    section_content.append(f"- {text};")
            
            if section_content:  # Если есть содержимое
                content.append(create_section_header(sections_translation[section], "_"))
                content.append("")
                content.extend(section_content)
                content.append("")
                content.append("_" * 70)
    
    return content

def save_as_txt(file_path, content):
    """Сохраняет в текстовый файл"""
    with open(file_path, 'w', encoding='utf-8') as file:
        file.write('\n'.join(content) if isinstance(content, list) else content)

def save_as_docx(file_path, content):
    """Сохраняет в формат Word с сохранением форматирования"""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Устанавливаем моноширинный шрифт и поля страницы
    style = doc.styles['Normal']
    style.font.name = 'Consolas'
    style.font.size = Pt(11)
    
    # Устанавливаем отступы страницы
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Добавляем каждую строку с сохранением форматирования
    for line in (content if isinstance(content, list) else content.split('\n')):
        # Пропускаем пустые строки
        if not line.strip():
            continue
            
        paragraph = doc.add_paragraph()
        
        # Обрабатываем разделители и заголовки
        if all(c == '=' for c in line.strip()) or all(c == '_' for c in line.strip()):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(line.strip())
        elif (':' in line) and ('=' in line or '_' in line):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(line.strip())
        else:
            # Для обычного текста
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Убираем лишний отступ в начале первой строки
            paragraph.paragraph_format.first_line_indent = Inches(0)
            
            # Обрабатываем строки с маркерами списка
            if line.lstrip().startswith('- '):
                paragraph.paragraph_format.left_indent = Inches(0.2)
                run = paragraph.add_run(line.lstrip())
            else:
                run = paragraph.add_run(line)
        
        run.font.name = 'Consolas'
    
    doc.save(file_path)

def save_as_pdf(file_path, content):
    """Сохраняет в формат PDF с поддержкой русского языка и форматированием"""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.units import mm
    
    try:
        # Регистрируем шрифт Arial с поддержкой русского языка
        try:
            pdfmetrics.registerFont(TTFont('Arial', 'C:/Windows/Fonts/arial.ttf'))
        except:
            try:
                pdfmetrics.registerFont(TTFont('Arial', '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf'))
            except:
                pass

        c = canvas.Canvas(file_path, pagesize=A4)
        width, height = A4
        margin = 50
        font_name = 'Arial'
        font_size = 11
        line_height = font_size * 1.5
        
        c.setFont(font_name, font_size)
        y = height - margin
        
        def add_text_line(text, y_pos):
            """Добавляет строку текста с учетом форматирования"""
            if not text.strip():
                return y_pos - line_height
            
            # Определяем максимальное количество символов в строке
            max_width = width - 2 * margin
            
            # Если строка начинается с маркера списка, добавляем отступ
            if text.lstrip().startswith('- '):
                margin_left = margin + 20
                max_width = width - margin - margin_left
            else:
                margin_left = margin
            
            # Разбиваем текст на строки
            lines = []
            current_line = []
            words = text.split()
            
            for word in words:
                test_line = ' '.join(current_line + [word])
                if c.stringWidth(test_line, font_name, font_size) <= max_width:
                    current_line.append(word)
                else:
                    if current_line:
                        lines.append(' '.join(current_line))
                        current_line = [word]
                    else:
                        lines.append(word)
            
            if current_line:
                lines.append(' '.join(current_line))
            
            # Отрисовываем каждую строку
            for line in lines:
                if y_pos < margin:  # Если достигли конца страницы
                    c.showPage()
                    c.setFont(font_name, font_size)
                    y_pos = height - margin
                
                c.drawString(margin_left, y_pos, line)
                y_pos -= line_height
            
            return y_pos
        
        # Обрабатываем каждую строку контента
        for line in (content if isinstance(content, list) else content.split('\n')):
            y = add_text_line(line, y)
        
        c.save()
    except Exception as e:
        raise Exception(f"Ошибка при сохранении PDF: {str(e)}")

def save_as_rtf(file_path, content):
    """Сохраняет в формат RTF с сохранением форматирования"""
    rtf_header = (
        '{\\rtf1\\ansi\\ansicpg1251\\deff0\\deflang1049\n'
        '{\\fonttbl{\\f0\\fmodern\\fcharset204 Consolas;}}\n'
        '\\viewkind4\\uc1\\pard\\f0\\fs22\n'
        '\\margl1440\\margr1440\n'
    )
    
    try:
        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(rtf_header)
            
            for line in (content if isinstance(content, list) else content.split('\n')):
                if not line.strip():
                    file.write('\\par\n')
                    continue
                
                # Определяем тип строки
                is_separator = all(c == '=' for c in line.strip()) or all(c == '_' for c in line.strip())
                is_header = (':' in line) and ('=' in line or '_' in line)
                is_metadata = any(prefix in line for prefix in ['Версия:', 'Дата:', 'Заголовок:'])
                
                # Настраиваем форматирование
                if is_separator or is_header:
                    file.write('\\ql\\li0 ')  # Без отступа
                elif is_metadata:
                    file.write('\\ql\\li0 ')  # Метаданные тоже без отступа
                else:
                    file.write('\\ql\\li0 ')  # Обычный текст без отступа
                
                # Обрабатываем строки с маркерами списка
                if line.lstrip().startswith('- '):
                    file.write('\\li360 ')  # Отступ для элементов списка
                
                # Экранируем специальные символы RTF
                line = line.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
                
                # Разбиваем строки с точкой с запятой на отдельные пункты
                if line.lstrip().startswith('- ') and ';' in line:
                    items = line.split(';')
                    for item in items:
                        if item.strip():
                            if not item.strip().startswith('-'):
                                item = '- ' + item.strip()
                            
                            # Конвертируем русские символы в RTF
                            rtf_line = ''
                            for char in item:
                                if ord(char) > 127:
                                    rtf_line += f'\\u{ord(char)}?'
                                else:
                                    rtf_line += char
                            
                            file.write(f'{rtf_line}\\par\n')
                else:
                    # Конвертируем русские символы в RTF
                    rtf_line = ''
                    for char in line:
                        if ord(char) > 127:
                            rtf_line += f'\\u{ord(char)}?'
                        else:
                            rtf_line += char
                    
                    # Добавляем перенос строки без дополнительных пробелов
                    file.write(f'{rtf_line}\\par\n')
            
            file.write('}')
    except Exception as e:
        raise Exception(f"Ошибка при сохранении RTF: {str(e)}")

def clear_editor_mode():
    """Очищает все поля в Editor Mode, кроме пути к папке и названия проекта."""
    if current_mode == "editor":
        form2_text.delete('1.0', tk.END)  # Основное текстовое поле
        version_entry.delete(0, tk.END)
        date_entry.delete(0, tk.END)
        description_entry.delete(0, tk.END)
        form4_text.delete('1.0', tk.END)  # Улучшения
        form6_text.delete('1.0', tk.END)  # Fixed Bugs
        form7_text.delete('1.0', tk.END)  # Notes
        
        # Сброс чекбоксов
        form6_check_var.set(False)
        form7_check_var.set(False)
        
def clear_release_mode():
    """Очищает все поля в Release Mode, кроме пути к папке и названия проекта."""
    if current_mode == "release":
        # Очищаем базовые поля
        version_entry.delete(0, tk.END)
        date_entry.delete(0, tk.END)
        title_entry.delete(0, tk.END)
        
        # Очищаем Summary
        summary_text.delete('1.0', tk.END)
        
        # Очищаем все секции с чекбоксами
        for section_data in section_widgets.values():
            # Очищаем все поля ввода в секции
            for entry in section_data["entries"]:
                entry.delete(0, tk.END)
            # Снимаем галочку с чекбокса
            section_data["check_var"].set(False)

def create_add_entry(entries_frame, entries, section_name):
    """Создает функцию добавления нового поля для конкретной секции"""
    def add_entry():
        entry = tk.Entry(entries_frame, width=120)
        entry.pack(padx=5, pady=2, fill=tk.X, expand=True)
        entries.append(entry)
        # Обновляем section_widgets после добавления нового поля
        section_widgets[section_name]["entries"] = entries
    return add_entry

def create_remove_entry(entries, section_name):
    """Создает функцию удаления поля для конкретной секции"""
    def remove_entry():
        if len(entries) > 1:  # Всегда оставляем хотя бы одно поле
            entry = entries.pop()
            entry.destroy()
            # Обновляем section_widgets после удаления поля
            section_widgets[section_name]["entries"] = entries
    return remove_entry

def create_menu():
    """Создает главное меню программы."""
    menu_bar = Menu(root)
    root.config(menu=menu_bar)
    
    # Меню "Файл"
    file_menu = Menu(menu_bar, tearoff=0)
    menu_bar.add_cascade(label="Файл", menu=file_menu)
    file_menu.add_command(label="Выход", command=root.quit)
    
    # Меню "Настройки"
    settings_menu = Menu(menu_bar, tearoff=0)
    menu_bar.add_cascade(label="Настройки", menu=settings_menu)
    settings_menu.add_command(label="Открыть настройки", command=open_settings)

# Обновляем функцию show_settings
def open_settings():
    """Показывает окно настроек."""
    settings_window = tk.Toplevel(root)
    settings_window.title("Настройки")
    
    # Делаем главное окно неактивным
    settings_window.grab_set()
    
    # Убираем стандартные кнопки окна
    settings_window.overrideredirect(True)
    
    # Создаем основную рамку для всего окна с более жирной границей
    main_frame = tk.Frame(settings_window, relief=tk.RIDGE, borderwidth=6)
    main_frame.pack(fill='both', expand=True)
    
    # Создаем заголовок окна
    title_frame = tk.Frame(main_frame, bg='#BBCCE3', cursor='hand2')  # Добавляем курсор руки
    title_frame.pack(fill='x')
    
    # Заголовок по центру
    title_label = tk.Label(title_frame, text="Настройки", bg='#BBCCE3', 
                          font=('Arial', 10, 'bold'), padx=10, pady=5,
                          cursor='hand2')  # Добавляем курсор руки
    title_label.pack(expand=True)
    
    # Линия-разделитель под заголовком
    separator = ttk.Separator(main_frame, orient='horizontal')
    separator.pack(fill='x')
    
    # Основное содержимое окна
    content_frame = tk.Frame(main_frame, padx=20, pady=20)
    content_frame.pack(fill='both', expand=True)
    
    # 1. Настройка режима по умолчанию
    default_mode_frame = tk.Frame(content_frame)
    default_mode_frame.pack(fill='x', pady=10)
    
    default_mode_label = tk.Label(default_mode_frame, text="Режим при запуске:")
    default_mode_label.pack(side='left', padx=5)
    
    default_mode_var = tk.StringVar(value=get_setting('default_mode', 'editor'))
    editor_radio = tk.Radiobutton(default_mode_frame, text="Editor Mode", 
                                 variable=default_mode_var, value='editor')
    release_radio = tk.Radiobutton(default_mode_frame, text="Release Mode", 
                                  variable=default_mode_var, value='release')
    editor_radio.pack(side='left', padx=5)
    release_radio.pack(side='left', padx=5)
    
    # 2. Настройка расположения версии в имени файла для Editor Mode
    version_pos_frame = tk.Frame(content_frame)
    version_pos_frame.pack(fill='x', pady=10)
    version_pos_label = tk.Label(version_pos_frame, 
                                text="для Editor Mode версия файла сохраняется:")
    version_pos_label.pack(side='left', padx=5)
    
    version_pos_var = tk.StringVar(value='end' if get_setting('version_at_end', True) else 'start')
    end_radio = tk.Radiobutton(version_pos_frame, text="В конце файла", 
                              variable=version_pos_var, value='end')
    start_radio = tk.Radiobutton(version_pos_frame, text="В начале файла", 
                                variable=version_pos_var, value='start')
    end_radio.pack(side='left', padx=5)
    start_radio.pack(side='left', padx=5)

    # 3. Настройка расположения версии в имени файла для Release Mode
    release_version_frame = tk.Frame(content_frame)
    release_version_frame.pack(fill='x', pady=10)
    release_version_label = tk.Label(release_version_frame, 
                                   text="для Release Mode версия файла сохраняется:")
    release_version_label.pack(side='left', padx=5)
    
    release_version_var = tk.StringVar(value='end' if get_setting('release_version_at_end', True) else 'start')
    release_end_radio = tk.Radiobutton(release_version_frame, text="В конце файла", 
                                      variable=release_version_var, value='end')
    release_start_radio = tk.Radiobutton(release_version_frame, text="В начале файла", 
                                        variable=release_version_var, value='start')
    release_end_radio.pack(side='left', padx=5)
    release_start_radio.pack(side='left', padx=5)

    # 4. Настройка формата файла
    file_format_frame = tk.Frame(content_frame)
    file_format_frame.pack(fill='x', pady=10)
    
    file_format_label = tk.Label(file_format_frame, 
                                text="Формат сохранения файла:")
    file_format_label.pack(side='left', padx=5)
    
    file_format_var = tk.StringVar(value=get_setting('default_file_format', '.txt'))
    formats = ['.txt', '.rtf', '.docx', '.pdf']
    
    # Создаем радиокнопки для каждого формата
    for format_type in formats:
        radio = tk.Radiobutton(file_format_frame, 
                              text=format_type,
                              variable=file_format_var,
                              value=format_type)
        radio.pack(side='left', padx=5)

    # 5. Настройка автосохранения
    autosave_frame = tk.Frame(content_frame)
    autosave_frame.pack(fill='x', pady=10)
    
    autosave_var = tk.BooleanVar(value=get_setting('autosave_enabled', False))
    autosave_check = ttk.Checkbutton(autosave_frame, 
                                    text="Сохранять введенный текст при закрытии программы",
                                    variable=autosave_var)
    autosave_check.pack(side='left', padx=5)
    
    # Кнопки управления
    button_frame = tk.Frame(main_frame)
    button_frame.pack(side='bottom', pady=10)
    
    def save_settings():
        save_setting('default_mode', default_mode_var.get())
        save_setting('version_at_end', version_pos_var.get() == 'end')
        save_setting('release_version_at_end', release_version_var.get() == 'end')
        save_setting('default_file_format', file_format_var.get())
        save_setting('autosave_enabled', autosave_var.get())
        
        # Если автосохранение включено, сохраняем текущие данные
        if autosave_var.get():
            save_form_data()
        
        settings_window.destroy()
    
    save_button = ttk.Button(button_frame, text="Сохранить", command=save_settings)
    close_button = ttk.Button(button_frame, text="Закрыть", 
                             command=settings_window.destroy)
    
    save_button.pack(side='left', padx=5)
    close_button.pack(side='left', padx=5)
    
    # Центрирование окна и настройка перемещения
    def center_settings_window():
        settings_window.update_idletasks()
        main_x = root.winfo_x()
        main_y = root.winfo_y()
        main_width = root.winfo_width()
        main_height = root.winfo_height()
        
        settings_width = settings_window.winfo_reqwidth()
        settings_height = settings_window.winfo_reqheight()
        
        x = main_x + (main_width - settings_width) // 2
        y = main_y + (main_height - settings_height) // 2
        
        settings_window.geometry(f"+{x}+{y}")
    
    # Делаем окно перемещаемым
    def start_move(event):
        settings_window.x = event.x
        settings_window.y = event.y

    def on_move(event):
        deltax = event.x - settings_window.x
        deltay = event.y - settings_window.y
        x = settings_window.winfo_x() + deltax
        y = settings_window.winfo_y() + deltay
        settings_window.geometry(f"+{x}+{y}")
    
    # Привязываем события перемещения
    title_frame.bind("<Button-1>", start_move)
    title_frame.bind("<B1-Motion>", on_move)
    title_label.bind("<Button-1>", start_move)
    title_label.bind("<B1-Motion>", on_move)
    
    settings_window.minsize(400, 250)
    center_settings_window()
    
    # Поднимаем окно поверх основного
    settings_window.lift()
    settings_window.focus_force()

def save_form_data(force_save=False):
    """Сохраняет данные всех форм"""
    form_data = get_setting('form_data', {})
    
    try:
        # Сохраняем данные Editor Mode
        if current_mode == "editor" and 'form1_entry' in globals():
            try:
                form_data['editor_mode'] = {
                    'project_name': form1_entry.get(),
                    'version': version_entry.get(),
                    'date': date_entry.get(),
                    'description': description_entry.get(),
                    'improvements': form4_text.get('1.0', tk.END),
                    'main_text': form2_text.get('1.0', tk.END),
                    'bugs': form6_text.get('1.0', tk.END),
                    'notes': form7_text.get('1.0', tk.END),
                    'bugs_enabled': form6_check_var.get(),
                    'notes_enabled': form7_check_var.get(),
                    'folder_path': folder_path_var.get()
                }
            except Exception as e:
                pass
        
        # Сохраняем данные Release Mode
        elif current_mode == "release" and 'release_project_entry' in globals():
            try:
                form_data['release_mode'] = {
                    'project_name': release_project_entry.get(),
                    'version': version_entry.get(),
                    'date': date_entry.get(),
                    'title': title_entry.get(),
                    'summary': summary_text.get('1.0', tk.END),
                    'summary_checked': summary_check_var.get(),
                    'folder_path': release_folder_path_var.get(),
                    'sections': {}
                }
                
                # Сохраняем данные всех секций, включая добавленные поля
                if 'section_widgets' in globals():
                    for section, widgets_data in section_widgets.items():
                        entries_data = []
                        for entry in widgets_data["entries"]:
                            try:
                                if entry.winfo_exists():
                                    entries_data.append(entry.get())
                            except tk.TclError:
                                continue
                        
                        # Всегда сохраняем все поля, даже пустые
                        form_data['release_mode']['sections'][section] = {
                            'entries': entries_data if entries_data else [""],
                            'checked': widgets_data["check_var"].get()
                        }
                        
                        # Обновляем также release_mode_data
                        release_mode_data["sections"][section] = {
                            'entries': entries_data if entries_data else [""],
                            'checked': widgets_data["check_var"].get()
                        }
            except Exception as e:
                pass
                
    except Exception as e:
        pass
        return

    # Всегда сохраняем в файл при force_save или включенном автосохранении
    if force_save or get_setting('autosave_enabled', False):
        save_setting('form_data', form_data)

def load_form_data():
    """Загружает сохраненные данные форм"""
    form_data = get_setting('form_data', {})
    
    # Всегда загружаем данные, независимо от настройки автосохранения
    if current_mode == "editor":
        if 'editor_mode' in form_data and 'form1_entry' in globals():
            try:
                ed_data = form_data['editor_mode']
                if 'folder_path' in ed_data and 'folder_path_var' in globals():
                    folder_path = ed_data['folder_path']
                    if folder_path and folder_path != "Выберите папку для сохранения проекта":
                        folder_path_var.set(folder_path)
                
                # Загружаем остальные данные Editor Mode
                if form1_entry.winfo_exists():
                    form1_entry.delete(0, tk.END)
                    version_entry.delete(0, tk.END)
                    date_entry.delete(0, tk.END)
                    description_entry.delete(0, tk.END)
                    form4_text.delete('1.0', tk.END)
                    form2_text.delete('1.0', tk.END)
                    form6_text.delete('1.0', tk.END)
                    form7_text.delete('1.0', tk.END)
                    
                    form1_entry.insert(0, ed_data.get('project_name', ''))
                    version_entry.insert(0, ed_data.get('version', ''))
                    date_entry.insert(0, ed_data.get('date', ''))
                    description_entry.insert(0, ed_data.get('description', ''))
                    form4_text.insert('1.0', ed_data.get('improvements', ''))
                    form2_text.insert('1.0', ed_data.get('main_text', ''))
                    form6_text.insert('1.0', ed_data.get('bugs', ''))
                    form7_text.insert('1.0', ed_data.get('notes', ''))
                    form6_check_var.set(ed_data.get('bugs_enabled', False))
                    form7_check_var.set(ed_data.get('notes_enabled', False))
            except (tk.TclError, AttributeError) as e:
                pass  # Удаляем print ошибки
    
    elif current_mode == "release":
        if 'release_mode' in form_data and 'release_project_entry' in globals():
            try:
                rel_data = form_data['release_mode']
                if 'folder_path' in rel_data and 'release_folder_path_var' in globals():
                    folder_path = rel_data['folder_path']
                    if folder_path and folder_path != "Выберите папку для сохранения проекта":
                        release_folder_path_var.set(folder_path)
                
                if release_project_entry.winfo_exists():
                    release_project_entry.delete(0, tk.END)
                    version_entry.delete(0, tk.END)
                    date_entry.delete(0, tk.END)
                    title_entry.delete(0, tk.END)
                    summary_text.delete('1.0', tk.END)
                    
                    release_project_entry.insert(0, rel_data.get('project_name', ''))
                    version_entry.insert(0, rel_data.get('version', ''))
                    date_entry.insert(0, rel_data.get('date', ''))
                    title_entry.insert(0, rel_data.get('title', ''))
                    summary_text.insert('1.0', rel_data.get('summary', ''))
                    summary_check_var.set(rel_data.get('summary_checked', False))
                    
                    # Загружаем данные секций
                    if 'sections' in rel_data and 'section_widgets' in globals():
                        for section, section_data in rel_data['sections'].items():
                            if section in section_widgets:
                                try:
                                    widgets_data = section_widgets[section]
                                    widgets_data['check_var'].set(section_data.get('checked', False))
                                    entries = section_data.get('entries', [])
                                    for entry, text in zip(widgets_data['entries'], entries):
                                        if entry.winfo_exists():
                                            entry.delete(0, tk.END)
                                            entry.insert(0, text)
                                except (tk.TclError, AttributeError) as e:
                                    pass  # Удаляем print ошибки
            except (tk.TclError, AttributeError) as e:
                pass  # Удаляем print ошибки

def on_closing():
    """Обработчик закрытия программы"""
    print("Закрытие программы...")
    
    # Всегда сохраняем текущее состояние перед закрытием
    save_form_data(force_save=True)
    
    # Если автосохранение выключено - очищаем файл
    if not get_setting('autosave_enabled', False):
        print("Автосохранение выключено, очищаем сохраненные данные...")
        save_setting('form_data', {})
    
    root.destroy()

# Функции для работы с настройками
def load_settings():
    """Загружает настройки из файла."""
    if os.path.exists(SETTINGS_FILE):
        try:
            with open(SETTINGS_FILE, 'r') as f:
                return json.load(f)
        except:
            return {}
    return {}

def save_setting(key, value):
    """Сохраняет настройку в файл."""
    settings = load_settings()
    settings[key] = value
    with open(SETTINGS_FILE, 'w') as f:
        json.dump(settings, f)

def get_setting(key, default=None):
    """Получает значение настройки"""
    settings = load_settings()
    value = settings.get(key, default)
    pass
    return value

# Create menu
menu_bar = Menu(root)
file_menu = Menu(menu_bar, tearoff=0)
file_menu.add_command(label="Editor Mode", command=switch_to_editor_mode)
file_menu.add_command(label="Release Mode", command=switch_to_release_mode)
file_menu.add_separator()
file_menu.add_command(label="Настройки", command=open_settings)
file_menu.add_command(label="Выход", command=exit_program)
menu_bar.add_cascade(label="Меню", menu=file_menu)
root.config(menu=menu_bar)

# Mode switch buttons
mode_switch_frame = tk.Frame(root)
editor_mode_btn = ttk.Button(mode_switch_frame, text="Editor Mode", command=switch_to_editor_mode)
release_mode_btn = ttk.Button(mode_switch_frame, text="Release Mode", command=switch_to_release_mode)
editor_mode_btn.pack(side=tk.LEFT, padx=5)
release_mode_btn.pack(side=tk.LEFT, padx=5)
mode_switch_frame.pack(pady=5)

# Main frame for dynamic content
main_frame = tk.Frame(root)
main_frame.pack(fill=tk.BOTH, expand=True)

# Editor Mode UI
def editor_mode_ui():
    """Создает интерфейс редактора"""
    global form1_entry, folder_path_var, form2_text, version_entry, date_entry
    global description_entry, form4_text, form6_text, form7_text
    global form6_check_var, form7_check_var
    global current_mode
    
    current_mode = "editor"
    clear_main_frame()
    
    folder_path_var = tk.StringVar(value=editor_mode_data["folder_path"])
    form6_check_var = tk.BooleanVar(value=editor_mode_data.get("form6_checked", False))
    form7_check_var = tk.BooleanVar(value=editor_mode_data.get("form7_checked", False))

    form1_frame = tk.Frame(main_frame)
    form1_label = tk.Label(form1_frame, text="Название проекта")
    form1_entry = tk.Entry(form1_frame, width=80)
    form1_entry.insert(0, editor_mode_data["project_name"])
    form1_label.pack(side=tk.LEFT, padx=5)
    form1_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
    form1_frame.pack(fill=tk.X, padx=10, pady=5)

    choose_folder_btn = ttk.Button(form1_frame, text="Выбрать папку", command=lambda: choose_folder(folder_path_var))
    choose_folder_btn.pack(side=tk.RIGHT, padx=5)
    folder_path_label = tk.Label(main_frame, textvariable=folder_path_var, anchor="w", relief=tk.SUNKEN)
    folder_path_label.pack(fill=tk.X, padx=10, pady=5)

    form2_frame = tk.Frame(main_frame)
    form2_label = tk.Label(form2_frame, text="Основное текстовое поле")
    form2_text = create_text_widget_with_menu(form2_frame, height=15)
    form2_text.insert("1.0", editor_mode_data["main_text"])
    form2_scroll = tk.Scrollbar(form2_frame, command=form2_text.yview)
    form2_text.configure(yscrollcommand=form2_scroll.set)
    form2_label.pack(anchor="w", padx=5)
    form2_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    form2_scroll.pack(side=tk.RIGHT, fill=tk.Y)
    form2_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

    form3_frame = tk.Frame(main_frame)
    
    # Версия
    version_frame = tk.Frame(form3_frame)
    version_label = tk.Label(version_frame, text="Версия", width=8, anchor="w")
    version_validate = root.register(validate_version)
    version_entry = tk.Entry(version_frame, width=12, validate="key", 
                           validatecommand=(version_validate, '%P'))
    version_entry.insert(0, editor_mode_data.get("version", ""))
    version_entry.bind('<KeyRelease>', format_version)
    version_label.pack(side=tk.LEFT, padx=2)
    version_entry.pack(side=tk.LEFT, padx=2)
    version_frame.pack(side=tk.LEFT)

    # Дата
    date_frame = tk.Frame(form3_frame)
    date_label = tk.Label(date_frame, text="Дата Релиза", width=10, anchor="w")
    date_validate = root.register(validate_date)
    date_entry = tk.Entry(date_frame, width=12, validate="key",
                         validatecommand=(date_validate, '%P'))
    date_entry.insert(0, editor_mode_data.get("date", ""))
    date_entry.bind('<KeyRelease>', format_date)
    date_label.pack(side=tk.LEFT, padx=2)
    date_entry.pack(side=tk.LEFT, padx=2)
    date_frame.pack(side=tk.LEFT)

    # Описание
    desc_frame = tk.Frame(form3_frame)
    desc_label = tk.Label(desc_frame, text="Описание", width=8, anchor="w")
    description_entry = tk.Entry(desc_frame, width=60)
    description_entry.insert(0, editor_mode_data.get("description", ""))
    desc_label.pack(side=tk.LEFT, padx=2)
    description_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=2)
    desc_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)

    form3_frame.pack(fill=tk.X, padx=10, pady=5)

    form4_frame = tk.Frame(main_frame)
    form4_label = tk.Label(form4_frame, text="Улучшения (Improvement)")
    form4_text, form4_scroll = create_text_widget_with_menu_and_format(form4_frame, height=7)
    form4_label.pack(anchor="w", padx=5)
    form4_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    form4_scroll.pack(side=tk.RIGHT, fill=tk.Y)
    form4_text.insert("1.0", editor_mode_data["improvements"])
    form4_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

    form6_frame = tk.Frame(main_frame)
    form6_check = tk.Checkbutton(form6_frame, text="Исправленные Ошибки (Fixed)", 
                                variable=form6_check_var)
    form6_text, form6_scroll = create_text_widget_with_menu_and_format(form6_frame, height=5)
    form6_check.pack(anchor="w", padx=5)
    form6_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    form6_scroll.pack(side=tk.RIGHT, fill=tk.Y)
    form6_text.insert("1.0", editor_mode_data["fixed_bugs"])
    form6_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
    

    form7_frame = tk.Frame(main_frame)
    form7_check = tk.Checkbutton(form7_frame, text="Заметки (Notes)", 
                                variable=form7_check_var)
    form7_text, form7_scroll = create_text_widget_with_menu_and_format(form7_frame, height=5)
    form7_check.pack(anchor="w", padx=5)
    form7_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    form7_scroll.pack(side=tk.RIGHT, fill=tk.Y)
    form7_text.insert("1.0", editor_mode_data["notes"])
    form7_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

    # Обновляем создание чекбоксов
    form6_check = tk.Checkbutton(form6_frame, text="Исправленные ошибки (Fixed)", 
                                variable=form6_check_var)
    form7_check = tk.Checkbutton(form7_frame, text="Заметки (Notes)", 
                                variable=form7_check_var)

    def _on_mousewheel(event, text_widget):
        try:
            text_widget.yview_scroll(int(-1*(event.delta/120)), "units")
        except tk.TclError:
            text_widget.unbind_all("<MouseWheel>")

    # Функция для привязки прокрутки к текстовому виджету
    def bind_text_scroll(text_widget):
        def _bind(event):
            text_widget.bind_all("<MouseWheel>", 
                lambda e, tw=text_widget: _on_mousewheel(e, tw))
        
        def _unbind(event):
            text_widget.unbind_all("<MouseWheel>")
        
        text_widget.bind("<Enter>", _bind)
        text_widget.bind("<Leave>", _unbind)

    # Привязываем прокрутку к каждому текстовому виджету
    for text_widget in [form2_text, form4_text, form6_text, form7_text]:
        bind_text_scroll(text_widget)   
     # Создаем фрейм для кнопок внизу
    bottom_buttons_frame = tk.Frame(main_frame)
    bottom_buttons_frame.pack(fill=tk.X, pady=10, padx=10)
    
    # Кнопка Clear EM слева
    clear_btn = ttk.Button(bottom_buttons_frame, text="Clear EM", command=clear_editor_mode)
    clear_btn.pack(side=tk.LEFT, padx=10)
    
    # Кнопка Создать справа
    create_btn = ttk.Button(bottom_buttons_frame, text="Создать", command=save_to_file)
    create_btn.pack(side=tk.RIGHT, padx=10)

    if get_setting('autosave_enabled', False):
        load_form_data()  # Загружаем данные после создания всех виджетов

# Release Mode UI
def release_mode_ui():
    """Создает интерфейс релиза"""
    global release_folder_path_var, release_project_entry, version_entry, date_entry
    global title_entry, summary_text, section_widgets, summary_check_var, canvas
    global current_mode
    
    current_mode = "release"
    clear_main_frame()
    
    # Инициализируем словарь для хранения виджетов секций
    section_widgets = {}

    # Инициализируем структуру данных, если она не существует
    if "sections" not in release_mode_data:
        release_mode_data["sections"] = {}
        for section in ["New Features", "Improvements", "Bug Fixes", "Known Issues", "Notes", "Feedback"]:
            release_mode_data["sections"][section] = {
                "entries": [""],  # Начинаем с одного пустого поля
                "checked": False
            }

    release_folder_path_var = tk.StringVar(value=release_mode_data["folder_path"])

    
    # Название проекта
    project_frame = tk.Frame(main_frame)
    project_label = tk.Label(project_frame, text="Название проекта")
    release_project_entry = tk.Entry(project_frame, width=80)
    release_project_entry.insert(0, release_mode_data.get("project_name", ""))
    project_label.pack(side=tk.LEFT, padx=5)
    release_project_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
    project_frame.pack(fill=tk.X, padx=10, pady=5)

    # Путь к папке
    folder_frame = tk.Frame(main_frame)
    folder_path_label = tk.Label(folder_frame, text="Путь к папке (Release Mode):")
    folder_path_display = tk.Label(folder_frame, textvariable=release_folder_path_var, 
                                 anchor="w", relief=tk.SUNKEN)
    choose_folder_btn = ttk.Button(folder_frame, text="Выбрать папку", 
                                 command=lambda: choose_folder(release_folder_path_var))
    folder_path_label.pack(side=tk.LEFT, padx=5)
    folder_path_display.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
    choose_folder_btn.pack(side=tk.RIGHT, padx=5)
    folder_frame.pack(fill=tk.X, padx=10, pady=5)

    # Создаем canvas и scrollbar
    canvas_frame = tk.Frame(main_frame)
    canvas = tk.Canvas(canvas_frame, width=800)  # Задаем начальную ширину
    scrollbar = tk.Scrollbar(canvas_frame, orient="vertical", command=canvas.yview)
    scrollable_frame = tk.Frame(canvas)

    # Настройка прокрутки
    def _on_frame_configure(event):
        canvas.configure(scrollregion=canvas.bbox("all"))

    def _on_mousewheel(event):
        # Проверяем, можно ли прокручивать в запрошенном направлении
        if event.delta > 0:  # Прокрутка вверх
            if canvas.yview()[0] <= 0:  # Если уже в самом верху
                return
        else:  # Прокрутка вниз
            if canvas.yview()[1] >= 1:  # Если уже в самом низу
                return
        canvas.yview_scroll(int(-1*(event.delta/120)), "units")

    # Привязываем прокрутку только к canvas_frame
    canvas_frame.bind("<Enter>", lambda e: canvas_frame.bind_all("<MouseWheel>", _on_mousewheel))
    canvas_frame.bind("<Leave>", lambda e: canvas_frame.unbind_all("<MouseWheel>"))

    # Настраиваем canvas
    canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
    scrollable_frame.bind("<Configure>", _on_frame_configure)
    canvas.configure(yscrollcommand=scrollbar.set)

    # Version, Date, Title 
    fields_frame = tk.Frame(main_frame)
    fields_frame.columnconfigure(5, weight=1)  # Даем больший вес колонке с Title
    
    # Version (формат XX.XX.XXX)
    version_label = tk.Label(fields_frame, text="Version", width=7, anchor="w")
    version_entry = tk.Entry(fields_frame, width=12)  # Ширина под формат XX.XX.XXX
    
    def validate_version(P):
        """Проверяет ввод версии в формате XX.XX.XXX"""
        if P == "":
            return True
        # Убираем 'v' если он есть в начале
        if P.startswith('v'):
            P = P[1:]
        parts = P.split('.')
        if len(parts) > 3:
            return False
        try:
            for i, part in enumerate(parts):
                if not part.isdigit() and part != "":
                    return False
                if part != "":
                    if i == 0 and len(part) > 2:  # Major Version (XX)
                        return False
                    if i == 1 and len(part) > 2:  # Minor Version (XX)
                        return False
                    if i == 2 and len(part) > 3:  # Patch Version (XXX)
                        return False
            return True
        except ValueError:
            return False
    
    def format_version(event):
        """Форматирует версию, добавляя 'v' и сохраняя ведущие нули."""
        entry = event.widget
        version = entry.get()
        
        if version:
            # Убираем 'v' если он есть
            if version.startswith('v'):
                version = version[1:]
                
            parts = version.split('.')
            formatted_parts = []
            for part in parts:
                if part:
                    formatted_parts.append(part)
                else:
                    formatted_parts.append(part)
            
            formatted_version = 'v' + '.'.join(formatted_parts)
            if formatted_version != 'v' + version:
                entry.delete(0, tk.END)
                entry.insert(0, formatted_version)
    
    version_validate = (main_frame.register(validate_version), '%P')
    version_entry.config(validate="key", validatecommand=version_validate)
    version_entry.bind('<FocusOut>', format_version)
    
    if release_mode_data.get("version"):
        version_entry.insert(0, release_mode_data["version"])
    
    version_label.grid(row=0, column=0, padx=(5,0), pady=2)
    version_entry.grid(row=0, column=1, padx=(2,5), pady=2, sticky="w")

    # Date
    date_label = tk.Label(fields_frame, text="Date", width=5, anchor="w")
    date_entry = tk.Entry(fields_frame, width=10)
    
    def validate_date(P):
        """Проверяет ввод даты в формате ДД.ММ.ГГГГ"""
        if P == "":
            return True
        
        if not all(c.isdigit() or c == '.' for c in P):
            return False
        
        parts = P.split('.')
        if len(parts) > 3:
            return False
        
        try:
            if len(parts) >= 1 and len(parts[0]) > 2:  # ДД
                return False
            if len(parts) >= 2 and len(parts[1]) > 2:  # ММ
                return False
            if len(parts) >= 3 and len(parts[2]) > 4:  # ГГГГ
                return False
                
            if len(parts[0]) == 2 and int(parts[0]) > 0:
                if not (1 <= int(parts[0]) <= 31):
                    return False
            if len(parts) > 1 and len(parts[1]) == 2 and int(parts[1]) > 0:
                if not (1 <= int(parts[1]) <= 12):
                    return False
            if len(parts) > 2 and len(parts[2]) == 4:
                if not (1000 <= int(parts[2]) <= 9999):
                    return False
            
            return True
        except (ValueError, IndexError):
            return True
    
    def format_date(event):
        """Форматирует дату, автоматически добавляя точки."""
        if event.char == '.':
            return
            
        entry = event.widget
        date = entry.get()
        clean_date = date.replace('.', '')
        formatted = ''
        
        if len(clean_date) > 0:
            formatted += clean_date[:2]
            if len(clean_date) > 2:
                formatted += '.' + clean_date[2:4]
                if len(clean_date) > 4:
                    formatted += '.' + clean_date[4:8]
        
        if formatted != date:
            entry.delete(0, tk.END)
            entry.insert(0, formatted)
    
    date_validate = (main_frame.register(validate_date), '%P')
    date_entry.config(validate="key", validatecommand=date_validate)
    date_entry.bind('<KeyRelease>', format_date)
    
    if release_mode_data.get("date"):
        date_entry.insert(0, release_mode_data["date"])
    
    date_label.grid(row=0, column=2, padx=(5,0), pady=2)
    date_entry.grid(row=0, column=3, padx=(2,5), pady=2, sticky="w")

    # Title
    title_label = tk.Label(fields_frame, text="Title", width=5, anchor="w")
    title_entry = tk.Entry(fields_frame)
    
    if release_mode_data.get("title"):
        title_entry.insert(0, release_mode_data["title"])
    
    title_label.grid(row=0, column=4, padx=(5,0), pady=2)
    title_entry.grid(row=0, column=5, padx=(2,5), pady=2, sticky="ew")

    fields_frame.pack(fill=tk.X, padx=10, pady=5)


    # Summary с чекбоксом и расширенными возможностями
    summary_frame = tk.Frame(main_frame)
    summary_check_var = tk.BooleanVar(value=release_mode_data.get("summary_checked", False))
    summary_check = tk.Checkbutton(summary_frame, text="Summary", variable=summary_check_var)
    summary_check.pack(anchor="w", padx=5)
    
    summary_text_frame = tk.Frame(summary_frame)
    summary_text, summary_scroll = create_text_widget_with_menu_and_format(summary_text_frame, height=7)
    summary_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    summary_scroll.pack(side=tk.RIGHT, fill=tk.Y)
    summary_text_frame.pack(fill=tk.BOTH, expand=True, padx=5)  # Текстовое поле всегда видимо
    
    # Восстанавливаем сохраненный текст
    if "summary" in release_mode_data:
        summary_text.delete("1.0", tk.END)
        summary_text.insert("1.0", release_mode_data.get("summary", ""))
    
    summary_frame.pack(fill=tk.X, padx=10, pady=5)
    
    # Вставляем сохраненный текст, если есть
    if "summary" in release_mode_data:
        summary_text.insert("1.0", release_mode_data.get("summary", ""))


 # Добавляем секции только один раз
    sections = ["New Features", "Improvements", "Bug Fixes", "Known Issues", "Notes", "Feedback"]
    for section in sections:
        section_frame = tk.Frame(scrollable_frame)
        entries_frame = tk.Frame(section_frame)
        entries = []
        
        # Получаем сохраненные данные для текущей секции
        section_data = release_mode_data["sections"].get(section, {"entries": [""], "checked": False})
        saved_entries = section_data.get("entries", [""])

        # Если нет сохраненных записей или они пустые, создаем одно пустое поле
        if not saved_entries or (len(saved_entries) == 1 and not saved_entries[0]):
            saved_entries = [""]
            
        # Создаем чекбокс
        check_var = tk.BooleanVar(value=section_data.get("checked", False))
        check = tk.Checkbutton(section_frame, text=section, variable=check_var)
        check.pack(anchor="w", padx=5)
        
        # Создаем контейнер для поля ввода и кнопок
        content_frame = tk.Frame(section_frame)
        
        # Фрейм для полей ввода
        entries_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)

        # Создаем поля ввода
        for saved_value in saved_entries:
            entry = tk.Entry(entries_frame, width=120)
            entry.pack(padx=5, pady=2, fill=tk.X, expand=True)
            if saved_value:  # Вставляем значение только если оно не пустое
                entry.insert(0, saved_value)
            entries.append(entry)
                
        # Создаем уникальные функции для текущей секции
        add_entry_func = create_add_entry(entries_frame, entries, section)
        remove_entry_func = create_remove_entry(entries, section)
        
        # Кнопки управления справа
        buttons_frame = tk.Frame(content_frame)
        add_btn = ttk.Button(buttons_frame, text="+", width=3, command=add_entry_func)
        remove_btn = ttk.Button(buttons_frame, text="-", width=3, command=remove_entry_func)
        
        buttons_frame.pack(side=tk.RIGHT, padx=5)
        add_btn.pack(side=tk.LEFT, padx=2)
        remove_btn.pack(side=tk.LEFT, padx=2)
        
        content_frame.pack(fill=tk.X, expand=True, padx=5)
        section_frame.pack(fill=tk.X, pady=5)
        
        # Сохраняем виджеты для последующего доступа
        section_widgets[section] = {
            "frame": section_frame,
            "entries": entries,
            "check_var": check_var
        }

    # Размещаем canvas и scrollbar
    canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
    canvas_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

    # Добавляем кнопки внизу
    bottom_buttons_frame = tk.Frame(main_frame)
    clear_btn = ttk.Button(bottom_buttons_frame, text="Clear RM", 
                          command=clear_release_mode)
    create_btn = ttk.Button(bottom_buttons_frame, text="Создать", 
                           command=save_to_file)
    
    clear_btn.pack(side=tk.LEFT, padx=10)
    create_btn.pack(side=tk.RIGHT, padx=10)
    bottom_buttons_frame.pack(fill=tk.X, pady=10)

    if get_setting('autosave_enabled', False):
        load_form_data()  # Загружаем данные после создания всех виджетов

if __name__ == "__main__":
    # Инициализация программы
    main()
    
    # Добавляем обработчик закрытия окна (Спасибо что досмотрели до конца этот сумасбродный код :)
    root.protocol("WM_DELETE_WINDOW", on_closing)
    
    # Запускаем главный цикл
    root.mainloop()